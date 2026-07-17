import io
import json
from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse
from database import get_db, get_user_tasks
from middleware.auth import get_current_user

router = APIRouter()

@router.get("/txt/{task_id}")
def export_txt(task_id: int, user=Depends(get_current_user)):
    conn = get_db()
    try:
        task = conn.execute("SELECT * FROM tasks WHERE id = ? AND user_id = ?", (task_id, user["id"])).fetchone()
    finally:
        conn.close()
    if not task:
        raise HTTPException(404, "Tarefa não encontrada")
    t = dict(task)
    content = f"NexusAgent AI - Resultado da Tarefa\n{'='*50}\n"
    content += f"Título: {t['title']}\nTipo: {t['task_type']}\nStatus: {t['status']}\n"
    content += f"Provedor: {t.get('ai_provider_used', 'N/A')}\nModelo: {t.get('ai_model_used', 'N/A')}\n"
    content += f"Criado: {t['created_at']}\n\n"
    content += f"DESCRIÇÃO:\n{t.get('description', '')}\n\n"
    content += f"RESULTADO:\n{t.get('output_data', 'Sem resultado')}\n"
    if t.get("error_message"):
        content += f"\nERRO:\n{t['error_message']}\n"
    return StreamingResponse(io.BytesIO(content.encode("utf-8")), media_type="text/plain", headers={"Content-Disposition": f"attachment; filename=tarefa_{task_id}.txt"})

@router.get("/pdf/{task_id}")
def export_pdf(task_id: int, user=Depends(get_current_user)):
    conn = get_db()
    try:
        task = conn.execute("SELECT * FROM tasks WHERE id = ? AND user_id = ?", (task_id, user["id"])).fetchone()
    finally:
        conn.close()
    if not task:
        raise HTTPException(404, "Tarefa não encontrada")
    t = dict(task)
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        y = height - 50
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, y, "NexusAgent AI - Tarefa")
        y -= 30
        c.setFont("Helvetica", 10)
        for line in [f"Título: {t['title']}", f"Tipo: {t['task_type']}", f"Status: {t['status']}", f"Provedor: {t.get('ai_provider_used', 'N/A')}", f"Modelo: {t.get('ai_model_used', 'N/A')}", f"Criado: {t['created_at']}", ""]:
            c.drawString(50, y, line)
            y -= 15
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, "Descrição:")
        y -= 15
        c.setFont("Helvetica", 10)
        for line in (t.get("description", "") or "").split("\n")[:20]:
            c.drawString(50, y, line[:90])
            y -= 12
        y -= 10
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, "Resultado:")
        y -= 15
        c.setFont("Helvetica", 10)
        for line in (t.get("output_data", "") or "Sem resultado").split("\n")[:40]:
            c.drawString(50, y, line[:90])
            y -= 12
            if y < 50:
                c.showPage()
                y = height - 50
                c.setFont("Helvetica", 10)
        c.save()
        buffer.seek(0)
        return StreamingResponse(buffer, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename=tarefa_{task_id}.pdf"})
    except ImportError:
        return StreamingResponse(io.BytesIO(b"ReportLab not installed"), status_code=501)

@router.get("/docx/{task_id}")
def export_docx(task_id: int, user=Depends(get_current_user)):
    conn = get_db()
    try:
        task = conn.execute("SELECT * FROM tasks WHERE id = ? AND user_id = ?", (task_id, user["id"])).fetchone()
    finally:
        conn.close()
    if not task:
        raise HTTPException(404, "Tarefa não encontrada")
    t = dict(task)
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.add_heading("NexusAgent AI - Tarefa", 0)
        doc.add_paragraph(f"Título: {t['title']}")
        doc.add_paragraph(f"Tipo: {t['task_type']}")
        doc.add_paragraph(f"Status: {t['status']}")
        doc.add_paragraph(f"Provedor: {t.get('ai_provider_used', 'N/A')}")
        doc.add_paragraph(f"Modelo: {t.get('ai_model_used', 'N/A')}")
        doc.add_paragraph("")
        doc.add_heading("Descrição", 1)
        doc.add_paragraph(t.get("description", ""))
        doc.add_heading("Resultado", 1)
        doc.add_paragraph(t.get("output_data", "Sem resultado"))
        if t.get("error_message"):
            doc.add_heading("Erro", 1)
            doc.add_paragraph(t["error_message"])
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=tarefa_{task_id}.docx"})
    except ImportError:
        return StreamingResponse(io.BytesIO(b"python-docx not installed"), status_code=501)

@router.get("/all-txt")
def export_all_txt(user=Depends(get_current_user)):
    tasks = get_user_tasks(user["id"], 200)
    content = "NexusAgent AI - Exportação Completa\n" + "=" * 60 + "\n\n"
    for t in tasks:
        content += f"[{t['id']}] {t['title']} ({t['task_type']}) - {t['status']}\n"
        content += f"  Criado: {t['created_at']}\n"
        content += f"  Resultado: {(t.get('output_data') or 'N/A')[:200]}\n\n"
    return StreamingResponse(io.BytesIO(content.encode("utf-8")), media_type="text/plain", headers={"Content-Disposition": "attachment; filename=nexusagent_tarefas.txt"})
